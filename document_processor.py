"""
Document Processing Module for Resume Analysis
Handles text extraction from various document formats: PDF, DOCX, DOC, TXT
"""

import os
import re
import logging
from typing import Optional, Dict, List, Tuple

# Document processing libraries
try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None
    PdfReader = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Universal document processor that can extract text from various file formats
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.doc': self._extract_doc_text,
            '.txt': self._extract_txt_text
        }
    
    def extract_text(self, file_path: str) -> Tuple[str, bool]:
        """
        Extract text from any supported document format
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            Tuple[str, bool]: (extracted_text, success_flag)
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return "", False
        
        # Get file extension
        _, ext = os.path.splitext(file_path.lower())
        
        if ext not in self.supported_formats:
            logger.error(f"Unsupported file format: {ext}")
            return "", False
        
        try:
            # Call appropriate extraction method
            text = self.supported_formats[ext](file_path)
            if text:
                # Clean and normalize the text
                cleaned_text = self._clean_text(text)
                return cleaned_text, True
            else:
                logger.warning(f"No text extracted from {file_path}")
                return "", False
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return "", False
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files using multiple methods"""
        text = ""
        
        # Method 1: Try pdfplumber (best for complex PDFs)
        if pdfplumber:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if text.strip():
                    logger.info(f"Successfully extracted text using pdfplumber from {file_path}")
                    return text
            except Exception as e:
                logger.warning(f"pdfplumber failed for {file_path}: {str(e)}")
        
        # Method 2: Try PyPDF2 (fallback)
        if PyPDF2 and PdfReader:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if text.strip():
                    logger.info(f"Successfully extracted text using PyPDF2 from {file_path}")
                    return text
            except Exception as e:
                logger.warning(f"PyPDF2 failed for {file_path}: {str(e)}")
        
        logger.error(f"All PDF extraction methods failed for {file_path}")
        return ""
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        if not Document:
            logger.error("python-docx not available for DOCX processing")
            return ""
        
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted text from DOCX: {file_path}")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {str(e)}")
            return ""
    
    def _extract_doc_text(self, file_path: str) -> str:
        """Extract text from DOC files (legacy format)"""
        # For DOC files, we'll try to read as text and clean up
        # This is a basic approach - for production, consider using python-docx2txt or antiword
        try:
            with open(file_path, 'rb') as file:
                content = file.read()
                # Try to decode and extract readable text
                text = content.decode('utf-8', errors='ignore')
                # Remove non-printable characters
                text = ''.join(char for char in text if char.isprintable() or char.isspace())
                
                logger.info(f"Extracted text from DOC file: {file_path}")
                return text
                
        except Exception as e:
            logger.error(f"Error extracting DOC text from {file_path}: {str(e)}")
            return ""
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT files"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                logger.info(f"Successfully read TXT file with UTF-8: {file_path}")
                return text
                
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                    logger.info(f"Successfully read TXT file with latin-1: {file_path}")
                    return text
                    
            except Exception as e:
                logger.error(f"Error reading TXT file {file_path}: {str(e)}")
                return ""
        
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {str(e)}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep alphanumeric and basic punctuation
        text = re.sub(r'[^\w\s\-\.\,\;\:\!\?\(\)\[\]\/\@\#\$\%\&\*\+\=]', ' ', text)
        
        # Remove excessive spaces again
        text = re.sub(r'\s+', ' ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_file_info(self, file_path: str) -> Dict[str, str]:
        """Get basic information about the file"""
        if not os.path.exists(file_path):
            return {"error": "File not found"}
        
        file_size = os.path.getsize(file_path)
        _, ext = os.path.splitext(file_path.lower())
        
        return {
            "file_path": file_path,
            "file_size": f"{file_size / 1024:.1f} KB",
            "file_extension": ext,
            "supported": ext in self.supported_formats
        }

# Global instance for easy import
document_processor = DocumentProcessor()
